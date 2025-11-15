#!/usr/bin/env python3
# utils.py — structure‑aware heading mapper (fix random H2/H3 placement)
#
# What changed in this build (vs the last one):
# 1) Replaced raw index‑by‑index style sequence overlay with a **structure‑aware mapper**.
#    We now detect *good heading candidates* in the refined text and assign the
#    exact number of H1/H2/H3 from the source to the best candidates **in order**.
#    This prevents headings from landing on arbitrary body paragraphs when your
#    refined draft has extra/shifted paragraph breaks.
# 2) Still honors explicit markdown markers in the refined text ('# ', '## ', '### ').
# 3) Keeps style name lookups (no style_id warnings) and the Arial 11/20/16/14 skeleton.
# 4) Reader remains robust (style name, outline level, char‑style, size/bold heuristics,
#    and traversal of body + tables).

from __future__ import annotations

import io
import os
import re
import pickle
import tempfile
import warnings
from pathlib import Path
from urllib.parse import urlparse, parse_qs
from typing import Dict, List, Tuple

import yaml
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from google.oauth2.credentials import Credentials
from google.oauth2 import service_account
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from googleapiclient.errors import HttpError
import time
import random
import os
import json as _json

# PDF and DOC support
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import docx2txt
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

# ---------------------------
# PDF and DOC extraction functions
# ---------------------------

def _extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    if not PDF_SUPPORT:
        raise ImportError("PyPDF2 is required for PDF support. Install with: pip install PyPDF2")
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")

def _extract_text_from_doc(file_path: str) -> str:
    """Extract text from DOC file"""
    if not DOC_SUPPORT:
        raise ImportError("python-docx2txt is required for DOC support. Install with: pip install python-docx2txt")
    
    try:
        import docx2txt
        text = docx2txt.process(file_path)
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOC: {str(e)}")

# ---------------------------
# Style helpers (name lookups)
# ---------------------------
_HEADING_ID_TO_NAME = {f"Heading{i}": f"Heading {i}" for i in range(1, 7)}
_HEADING_NAME_TO_NAME = {f"Heading {i}": f"Heading {i}" for i in range(1, 7)}
_CANON_STYLE = {"Normal": "Normal", **_HEADING_ID_TO_NAME, **_HEADING_NAME_TO_NAME}

def _canon_style_name(key: str) -> str:
    return _CANON_STYLE.get(key, key)

def _get_style_by_name(doc: Document, key: str):
    name = _canon_style_name(key)

    # First, try explicit name match (avoids deprecated style_id lookup)
    for st in doc.styles:
        try:
            if getattr(st, "name", None) == name:
                return st
        except Exception:
            continue

    # Fallback: silence docx's deprecated style_id lookup warning
    import warnings
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        return doc.styles[name]

    # If nothing works, raise explicit error
    raise KeyError(f"Style not found by name: {name}")

# ----------------
# Heuristics YAML
# ----------------

def load_heuristics(path: str = None) -> dict:
    """Load heuristics YAML. Resolves relative to backend/config if not found in CWD.

    Search order:
      1) Given absolute path (as-is)
      2) Relative path from current working directory
      3) Relative path from backend/config/ directory
    """
    if path is None:
        path = "heuristics.yaml"
    
    candidates = []
    if os.path.isabs(path):
        candidates.append(path)
    else:
        # CWD
        candidates.append(os.path.abspath(path))
        # Backend config directory (utils.py is in backend/, so dirname(__file__) gives backend/)
        backend_dir = os.path.dirname(os.path.abspath(__file__))
        config_path = os.path.join(backend_dir, 'config', path)
        candidates.append(config_path)

    for p in candidates:
        try:
            if os.path.exists(p):
                with open(p, 'r', encoding='utf-8') as f:
                    return yaml.safe_load(f)
        except Exception:
            continue
    raise FileNotFoundError(f"Heuristics YAML not found. Tried: {candidates}")

# ------------------------------
# History analysis (MVP profile)
# ------------------------------

def derive_history_profile(history_path: str = None, max_entries: int = 50) -> Dict[str, float]:
    """Compute a lightweight session profile from recent history and logs.

    Returns a dict with normalized hints in 0..1 range:
      - brevity_bias: tendency to shorten content
      - formality_bias: tendency to prefer formal tone
      - structure_bias: tendency to add headings/separators
    """
    if history_path is None:
        backend_dir = os.path.dirname(os.path.dirname(__file__))
        history_path = os.path.join(backend_dir, 'data', 'recent_history.json')
    
    profile = {"brevity_bias": 0.5, "formality_bias": 0.5, "structure_bias": 0.5}
    try:
        # Look at logs/refiner.log for PASS_TOGGLES signals
        backend_dir = os.path.dirname(os.path.dirname(__file__))
        log_path = os.path.join(backend_dir, 'logs', 'refiner.log')
        changes: List[float] = []
        redundants: List[int] = []
        with open(log_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()[-max_entries:]
        for ln in lines:
            if 'PASS_TOGGLES' in ln:
                # extract edits_per_100w and sentences vs prev
                try:
                    # crude parse of key=value tokens
                    tokens = dict(
                        (kv.split('=')[0], kv.split('=')[1])
                        for kv in (seg.strip() for seg in ln.split() if '=' in seg)
                    )
                    edits = float(tokens.get('edits_per_100w', '0'))
                    sents = float(tokens.get('sentences', '0'))
                    sents_prev = float(tokens.get('sentences_prev', '0'))
                    changes.append(edits)
                    if sents_prev > 0 and sents < sents_prev:
                        redundants.append(1)
                except Exception:
                    continue
        # Normalize simple tendencies
        if changes:
            # Higher edits/100w => higher brevity/structure pressure
            avg_edits = sum(changes) / max(1.0, float(len(changes)))
            profile["brevity_bias"] = max(0.0, min(1.0, avg_edits / 40.0))
            profile["structure_bias"] = max(0.0, min(1.0, avg_edits / 50.0))
        if redundants:
            profile["structure_bias"] = max(profile["structure_bias"], 0.6)

        # Peek recent_history.json for any user choices (placeholder hook)
        try:
            if not os.path.isabs(history_path):
                backend_dir = os.path.dirname(os.path.dirname(__file__))
                hist_full = os.path.join(backend_dir, 'data', os.path.basename(history_path))
            else:
                hist_full = history_path
            if os.path.exists(hist_full):
                with open(hist_full, 'r', encoding='utf-8') as hf:
                    _json.load(hf)
        except Exception:
            pass
    except Exception:
        pass
    return profile

# ----------------------
# Reader (robust H1‑H3)
# ----------------------

def read_text_from_file(file_path: str) -> str:
    # Guard: explicit existence check with clear error for callers/GUI
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Input file does not exist: {file_path}")
    lower = file_path.lower()
    if lower.endswith('.txt') or lower.endswith('.md'):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif lower.endswith('.pdf'):
        return _extract_text_from_pdf(file_path)
    elif lower.endswith('.doc'):
        return _extract_text_from_doc(file_path)
    elif lower.endswith('.docx'):
        pass  # Continue with existing DOCX logic
    else:
        raise ValueError(f"Unsupported file type: {file_path}")

    doc = Document(file_path)

    def run_size_pt(run) -> float | None:
        sz = getattr(getattr(run, 'font', None), 'size', None)
        if sz:
            try:
                return float(sz.pt)
            except Exception:
                pass
        try:
            rPr = run._element.rPr
            if rPr is not None and getattr(rPr, 'sz', None) is not None and rPr.sz.val is not None:
                return float(rPr.sz.val) / 2.0
        except Exception:
            pass
        return None

    def paragraph_heading_level(p) -> int:
        name = getattr(getattr(p, 'style', None), 'name', '') or ''
        m = re.search(r'heading\s*(\d)', name, flags=re.IGNORECASE)
        if m:
            try:
                return min(int(m.group(1)), 6)
            except ValueError:
                pass
        lname = name.lower()
        if lname == 'title':
            return 1
        if lname in ('subtitle', 'sub-title', 'sub title'):
            return 2
        try:
            ppr = p._element.pPr
            if ppr is not None and getattr(ppr, 'outlineLvl', None) is not None and ppr.outlineLvl.val is not None:
                return min(int(ppr.outlineLvl.val) + 1, 6)
        except Exception:
            pass
        for r in p.runs:
            try:
                rstyle = getattr(getattr(r, 'style', None), 'name', '') or ''
                m2 = re.search(r'heading\s*(\d)', rstyle, flags=re.IGNORECASE)
                if m2:
                    return min(int(m2.group(1)), 6)
            except Exception:
                pass
        sizes = []
        any_bold = False
        has_text = False
        for r in p.runs:
            if (r.text or '').strip():
                has_text = True
                s = run_size_pt(r)
                if s:
                    sizes.append(s)
                b = getattr(getattr(r, 'font', None), 'bold', None)
                any_bold = any_bold or bool(b) if b is not None else any_bold
        if has_text and sizes:
            import statistics
            try:
                size_pt = statistics.median(sizes)
            except Exception:
                size_pt = sizes[0]
            def approx(val, target, tol=0.6):
                return abs(val - target) <= tol
            if approx(size_pt, 20.0) and not any_bold:
                return 1
            if approx(size_pt, 16.0) and not any_bold:
                return 2
            if approx(size_pt, 14.0) and any_bold:
                return 3
        return 0

    lines: List[str] = []

    def walk(el) -> None:
        if hasattr(el, 'paragraphs'):
            for p in el.paragraphs:
                text = p.text or ''
                lvl = paragraph_heading_level(p)
                if lvl > 0:
                    lines.append('#' * lvl + ' ' + text)
                else:
                    lines.append(text)
        if hasattr(el, 'tables'):
            for t in el.tables:
                for row in t.rows:
                    for cell in row.cells:
                        walk(cell)

    walk(doc)
    return '\n'.join(lines)

# -------------------------
# Skeleton + seq from DOCX
# -------------------------

StyleRecord = Dict[str, object]
StyleSkeleton = Dict[str, StyleRecord]

_DEF_SKEL: StyleSkeleton = {
    "Normal":    {"font_name":"Arial","font_size_pt":11,"bold":False,"color_rgb":(0,0,0),
                  "space_before_pt":0,"space_after_pt":12,"line_spacing":1.0},
    "Heading 1": {"font_name":"Arial","font_size_pt":20,"bold":False,"color_rgb":(0,0,0),
                  "space_before_pt":12,"space_after_pt":12,"line_spacing":1.0},
    "Heading 2": {"font_name":"Arial","font_size_pt":16,"bold":False,"color_rgb":(0,0,0),
                  "space_before_pt":12,"space_after_pt":6,"line_spacing":1.0},
    "Heading 3": {"font_name":"Arial","font_size_pt":14,"bold":True ,"color_rgb":(0,0,0),
                  "space_before_pt":6,"space_after_pt":6,"line_spacing":1.0},
}

def _extract_style_record(doc: Document, style_name: str) -> StyleRecord:
    rec: StyleRecord = dict(_DEF_SKEL.get(style_name, {}))
    try:
        st = _get_style_by_name(doc, style_name)
    except KeyError:
        return rec
    if hasattr(st, 'font'):
        if st.font.name:
            rec['font_name'] = st.font.name
        if st.font.size:
            rec['font_size_pt'] = int(st.font.size.pt)
        if st.font.bold is not None:
            rec['bold'] = bool(st.font.bold)
        if st.font.color and st.font.color.rgb:
            rgb = st.font.color.rgb
            rec['color_rgb'] = (rgb[0], rgb[1], rgb[2])
    if hasattr(st, 'paragraph_format'):
        pf = st.paragraph_format
        if pf.space_before:
            rec['space_before_pt'] = int(pf.space_before.pt)
        if pf.space_after:
            rec['space_after_pt'] = int(pf.space_after.pt)
        if pf.line_spacing:
            try:
                rec['line_spacing'] = float(pf.line_spacing)
            except Exception:
                pass
    return rec

def make_style_skeleton_from_docx(source_docx_path: str) -> StyleSkeleton:
    doc = Document(source_docx_path)
    skel: StyleSkeleton = {k: dict(v) for k, v in _DEF_SKEL.items()}
    for sname in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        src = _extract_style_record(doc, sname)
        for k, v in src.items():
            if v is not None:
                skel[sname][k] = v
    return skel

def _apply_style_defaults(doc: Document, style_name: str, rec: StyleRecord) -> None:
    try:
        st = _get_style_by_name(doc, style_name)
    except KeyError:
        return
    st.font.name = rec['font_name']
    if hasattr(st, '_element') and hasattr(st._element, 'rPr') and st._element.rPr is not None:
        st._element.rPr.rFonts.set(qn('w:eastAsia'), rec['font_name'])
    from docx.shared import Pt as _Pt
    st.font.size = _Pt(rec['font_size_pt'])
    st.font.bold = bool(rec.get('bold', False))
    r, g, b = rec['color_rgb']
    st.font.color.rgb = RGBColor(int(r), int(g), int(b))
    pf = st.paragraph_format
    pf.space_before = _Pt(rec['space_before_pt'])
    pf.space_after  = _Pt(rec['space_after_pt'])
    pf.line_spacing = rec['line_spacing']

def _set_paragraph_format(p, rec: StyleRecord) -> None:
    from docx.shared import Pt as _Pt
    pf = p.paragraph_format
    pf.space_before = _Pt(rec['space_before_pt'])
    pf.space_after  = _Pt(rec['space_after_pt'])
    pf.line_spacing = rec['line_spacing']

# ---------------------------------------------
# Structure‑aware heading target selection
# ---------------------------------------------

def _heading_counts_from_seq(seq: List[str]) -> Tuple[int, int, int]:
    return (seq.count('Heading 1'), seq.count('Heading 2'), seq.count('Heading 3'))

_TITLEY_RE = re.compile(r"^(?:[A-Z][\w\-']+\s+){1,12}[A-Za-z0-9\-']{1,}$")
_PUNCTUATION_END = re.compile(r"[.!?…]$")


def _score_heading_candidate(line: str) -> float:
    """Higher score => more heading‑like.
    Heuristics: shortish line, not ending with period, title‑case ratio, no leading bullet.
    """
    s = line.strip()
    if not s:
        return 0.0
    if s[:2] in ("- ", "• ", "* "):
        return 0.0
    length = len(s)
    if length > 120:
        return 0.0
    # Avoid sentences that look like body text (end with period/question/ellipsis)
    if _PUNCTUATION_END.search(s):
        body_penalty = 0.25
    else:
        body_penalty = 0.0
    # Title‑case ratio
    words = [w for w in re.split(r"\s+", s) if w]
    caps = sum(1 for w in words if w[:1].isupper())
    ratio = (caps / max(1, len(words)))
    title_bonus = ratio * 0.6
    # Ultra short/long shaping
    length_bonus = 0.4 if 8 <= length <= 80 else 0.1 if length < 8 else 0.0
    return max(0.0, title_bonus + length_bonus - body_penalty)


def _pick_heading_targets(lines: List[str], want_h1: int, want_h2: int, want_h3: int) -> Dict[int, str]:
    """Return a map {line_index: style_name} choosing heading‑like lines in order.
    Priority:
      1) Explicit markers '# ', '## ', '### '
      2) Top‑scoring title‑like lines in reading order
    Count of each level matches source counts; extra markers beyond source caps are downgraded to Normal.
    """
    n = len(lines)
    assigned: Dict[int, str] = {}

    # 1) Collect explicit markers
    explicit_h1 = [i for i,l in enumerate(lines) if l.startswith('# ')]
    explicit_h2 = [i for i,l in enumerate(lines) if l.startswith('## ')]
    explicit_h3 = [i for i,l in enumerate(lines) if l.startswith('### ')]

    def take_first_k(idxs: List[int], k: int) -> List[int]:
        out = []
        for i in idxs:
            if len(out) >= k: break
            out.append(i)
        return out

    mark_h1 = take_first_k(explicit_h1, want_h1)
    mark_h2 = take_first_k(explicit_h2, want_h2)
    mark_h3 = take_first_k(explicit_h3, want_h3)

    for i in mark_h1: assigned[i] = 'Heading 1'
    for i in mark_h2: assigned[i] = 'Heading 2'
    for i in mark_h3: assigned[i] = 'Heading 3'

    # remaining counts to fill via candidates
    left_h1 = max(0, want_h1 - len(mark_h1))
    left_h2 = max(0, want_h2 - len(mark_h2))
    left_h3 = max(0, want_h3 - len(mark_h3))

    if left_h1 + left_h2 + left_h3 == 0:
        return assigned

    # 2) Score candidates (skip already assigned and blank lines)
    scored: List[Tuple[int, float]] = []
    for i, l in enumerate(lines):
        if i in assigned: 
            continue
        if not l.strip():
            continue
        # If line starts with a marker but we don't need more of that level, treat as body
        if l.startswith('#'):
            continue
        scored.append((i, _score_heading_candidate(l)))

    # Keep only reasonably title‑like lines
    scored = [(i,s) for (i,s) in scored if s >= 0.35]
    scored.sort(key=lambda t: (t[0], -t[1]))  # reading order, tie‑break by score

    # Assign in reading order: fill H1, then H2s, then H3s
    for i,_ in scored:
        if left_h1 > 0:
            assigned[i] = 'Heading 1'; left_h1 -= 1; continue
        if left_h2 > 0:
            assigned[i] = 'Heading 2'; left_h2 -= 1; continue
        if left_h3 > 0:
            assigned[i] = 'Heading 3'; left_h3 -= 1; continue
        if left_h1 + left_h2 + left_h3 == 0:
            break

    return assigned

# ----------------------------
# Writer (structure‑aware map)
# ----------------------------

def write_docx_with_skeleton(text: str, out_path: str, skel: Dict[str, Dict[str, object]], seq: List[str] | None) -> None:
    doc = Document()
    for s in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        _apply_style_defaults(doc, s, skel[s])

    lines = text.splitlines()

    # Derive desired counts from source sequence
    want_h1, want_h2, want_h3 = _heading_counts_from_seq(seq or [])

    # Build structure‑aware target map
    targets = _pick_heading_targets(lines, want_h1, want_h2, want_h3)

    for i, raw in enumerate(lines):
        line = raw if raw is not None else ""
        # Strip explicit markers, but they still influence assignment
        if line.startswith('### '):
            content = line[4:]
            style = targets.get(i, 'Heading 3' if want_h3 else 'Normal')
        elif line.startswith('## '):
            content = line[3:]
            style = targets.get(i, 'Heading 2' if want_h2 else 'Normal')
        elif line.startswith('# '):
            content = line[2:]
            style = targets.get(i, 'Heading 1' if want_h1 else 'Normal')
        else:
            content = line
            style = targets.get(i, 'Normal')

        p = doc.add_paragraph(content, style=style)
        _set_paragraph_format(p, skel[style if style in skel else 'Normal'])

    doc.save(out_path)

# ---------------------------
# write_text_to_file (DOCX)
# ---------------------------

def write_text_to_file(output_dir: str, base_name: str, ext: str, text: str,
                       original_file: str, iteration: int) -> str:
    # Normalize base name: avoid double pass suffixes like _pass2_pass1
    try:
        # Remove one or more trailing _pass<digits> groups (e.g., _pass2_pass3)
        base_clean = re.sub(r"(?:_pass\d+)+$", "", base_name, flags=re.IGNORECASE)
    except Exception:
        base_clean = base_name
    suffix = ext if ext.startswith('.') else f'.{ext}'
    filename = f"{base_clean}_pass{iteration}{suffix}"
    # Create an isolated, unique temp subdirectory per write to avoid collisions in batch runs
    # Keep stable basename for downstream move/upload while isolating the directory
    temp_dir = tempfile.mkdtemp(prefix=f"refiner_{base_name}_")
    local_path = os.path.join(temp_dir, filename)

    if suffix in ['.txt', '.md']:
        with open(local_path, 'w', encoding='utf-8') as f:
            f.write(text)
        return local_path

    # Learn skeleton and source sequence counts from the original DOCX (if present)
    if str(original_file).lower().endswith('.docx') and os.path.exists(original_file):
        try:
            skel = make_style_skeleton_from_docx(original_file)
        except Exception:
            skel = {k: dict(v) for k, v in _DEF_SKEL.items()}
        try:
            seq = make_style_sequence_from_docx(original_file)
        except Exception:
            seq = []
    else:
        skel = {k: dict(v) for k, v in _DEF_SKEL.items()}
        seq = []

    write_docx_with_skeleton(text, local_path, skel, seq)
    return local_path

# ------------------------
# Sequence from source DOCX
# ------------------------

def make_style_sequence_from_docx(source_docx_path: str) -> List[str]:
    doc = Document(source_docx_path)
    seq: List[str] = []

    def level_to_name(lvl: int) -> str:
        return {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}.get(lvl, 'Normal')

    def para_level(p) -> int:
        name = getattr(getattr(p, 'style', None), 'name', '') or ''
        m = re.search(r'heading\s*(\d)', name, flags=re.IGNORECASE)
        lvl = int(m.group(1)) if m else 0
        if lvl == 0:
            try:
                ppr = p._element.pPr
                if ppr is not None and getattr(ppr, 'outlineLvl', None) is not None and ppr.outlineLvl.val is not None:
                    lvl = int(ppr.outlineLvl.val) + 1
            except Exception:
                pass
        if lvl == 0:
            for r in p.runs:
                try:
                    rname = getattr(getattr(r, 'style', None), 'name', '') or ''
                    m2 = re.search(r'heading\s*(\d)', rname, flags=re.IGNORECASE)
                    if m2:
                        lvl = int(m2.group(1)); break
                except Exception:
                    pass
        return lvl if 1 <= lvl <= 3 else 0

    def walk(el):
        if hasattr(el, 'paragraphs'):
            for p in el.paragraphs:
                seq.append(level_to_name(para_level(p)))
        if hasattr(el, 'tables'):
            for t in el.tables:
                for row in t.rows:
                    for cell in row.cells:
                        walk(cell)

    walk(doc)
    return seq

# ------------------------------
# Drive + OAuth + Docs helpers
# ------------------------------

def extract_drive_file_id(link_or_id: str) -> str:
    m = re.search(r'/d/([A-Za-z0-9_-]+)', link_or_id)
    if m:
        return m.group(1)
    qs = parse_qs(urlparse(link_or_id).query)
    if 'id' in qs:
        return qs['id'][0]
    clean = link_or_id.strip()
    if re.fullmatch(r'[A-Za-z0-9_-]{10,}', clean):
        return clean
    raise ValueError(f"Could not parse Drive file ID from '{link_or_id}'")


def download_drive_file(link_or_id: str, dest_path: str) -> str:
    file_id = extract_drive_file_id(link_or_id)
    creds = get_google_credentials()
    drive = build('drive', 'v3', credentials=creds)
    # Retry with exponential backoff for transient and quota errors
    max_retries = int(os.getenv('API_MAX_RETRIES', '6'))
    backoff = float(os.getenv('API_BACKOFF_START', '1.0'))
    backoff_cap = float(os.getenv('API_BACKOFF_CAP', '8.0'))
    for attempt in range(max_retries):
        try:
            request = drive.files().get_media(fileId=file_id)
            fh = io.FileIO(dest_path, 'wb')
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            fh.close()
            return dest_path
        except HttpError as e:
            status = getattr(getattr(e, 'resp', None), 'status', None)
            if status == 404:
                raise FileNotFoundError(f"Drive file not found: {file_id}")
            if status in (429, 403, 500, 503):
                time.sleep(backoff + random.uniform(0, 0.25))
                backoff = min(backoff_cap, backoff * 2.0)
                continue
            raise

OAUTH_SCOPES = ['https://www.googleapis.com/auth/drive', 'https://www.googleapis.com/auth/documents']

def get_google_credentials(credentials_path: str = None, token_path: str = None) -> Credentials:
    # Check for service account credentials from environment variable first (for Vercel/serverless)
    google_creds_json = os.getenv('GOOGLE_CREDENTIALS_JSON')
    if google_creds_json:
        try:
            import tempfile
            # Parse JSON and write to temp file
            creds_data = json.loads(google_creds_json)
            is_vercel = os.getenv('VERCEL') == '1' or os.getenv('VERCEL_ENV') is not None
            if is_vercel:
                temp_dir = Path('/tmp/config')
                temp_dir.mkdir(parents=True, exist_ok=True)
                temp_file = temp_dir / 'google_credentials.json'
            else:
                temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False)
                temp_file_path = Path(temp_file.name)
                temp_file.close()
                temp_file = temp_file_path
            
            with open(temp_file, 'w') as f:
                json.dump(creds_data, f)
            
            creds = service_account.Credentials.from_service_account_file(
                str(temp_file),
                scopes=OAUTH_SCOPES
            )
            return creds
        except Exception as e:
            print(f"Warning: Failed to load Google credentials from env var: {e}")
    
    # Check for service account file first
    # Prefer explicit env var; default to backend/config/google_credentials.json
    backend_dir = os.path.dirname(os.path.dirname(__file__))
    default_service_account = os.path.join(backend_dir, 'config', 'google_credentials.json')
    service_account_file = os.getenv('GOOGLE_SERVICE_ACCOUNT_FILE', default_service_account)
    
    if credentials_path is None:
        credentials_path = os.path.join(backend_dir, 'config', 'credentials.json')
    if token_path is None:
        token_path = os.path.join(backend_dir, 'config', 'token.json')
    if os.path.exists(service_account_file):
        try:
            # Use service account authentication
            creds = service_account.Credentials.from_service_account_file(
                service_account_file,
                scopes=OAUTH_SCOPES
            )
            return creds
        except Exception as e:
            print(f"Warning: Failed to load service account credentials: {e}")
            print("Falling back to OAuth flow...")
    
    # Fall back to OAuth flow
    creds = None
    token_file = Path(token_path)
    if token_file.exists():
        creds = pickle.loads(token_file.read_bytes())
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(credentials_path, OAUTH_SCOPES)
            creds = flow.run_local_server(port=0)
        token_file.write_bytes(pickle.dumps(creds))
    return creds

def get_drive_service_oauth():
    """Get Google Drive service using OAuth credentials"""
    try:
        from googleapiclient.discovery import build
        creds = get_google_credentials()
        if not creds:
            return None
        service = build('drive', 'v3', credentials=creds)
        return service
    except Exception as e:
        import traceback
        print(f"Error building Drive service: {e}\n{traceback.format_exc()}")
        return None

# -----------------------------
# Google Docs round‑trip pieces
# -----------------------------

def fetch_gdoc_body(creds, document_id: str) -> list:
    service = build('docs', 'v1', credentials=creds)
    max_retries = int(os.getenv('API_MAX_RETRIES', '6'))
    backoff = float(os.getenv('API_BACKOFF_START', '1.0'))
    backoff_cap = float(os.getenv('API_BACKOFF_CAP', '8.0'))
    for attempt in range(max_retries):
        try:
            doc = service.documents().get(documentId=document_id).execute()
            return doc.get('body', {}).get('content', [])
        except HttpError as e:
            status = getattr(getattr(e, 'resp', None), 'status', None)
            if status in (429, 403, 500, 503):
                time.sleep(backoff + random.uniform(0, 0.25))
                backoff = min(backoff_cap, backoff * 2.0)
                continue
            raise


def write_gdoc_to_docx(creds, document_id: str, template_path: str, output_path: str) -> None:
    doc = Document(template_path)
    content = fetch_gdoc_body(creds, document_id)
    for element in content:
        para = element.get('paragraph')
        if para is None:
            continue
        text_runs = [run.get('textRun', {}).get('content', '') for run in para.get('elements', []) if run.get('textRun')]
        full_text = ''.join(text_runs).rstrip("\r\n")
        style_type = para.get('paragraphStyle', {}).get('namedStyleType', '')
        if   style_type == 'HEADING_1': doc.add_heading(full_text, level=1)
        elif style_type == 'HEADING_2': doc.add_heading(full_text, level=2)
        elif style_type == 'HEADING_3': doc.add_heading(full_text, level=3)
        elif style_type == 'HEADING_4': doc.add_heading(full_text, level=4)
        elif style_type == 'HEADING_5': doc.add_heading(full_text, level=5)
        elif style_type == 'HEADING_6': doc.add_heading(full_text, level=6)
        else:
            p = doc.add_paragraph(full_text)
            p.style = _get_style_by_name(doc, 'Normal')
    doc.save(output_path)

# Optional hook (no‑op)

def apply_named_styles(local_docx_path: str, doc_id: str, creds: Credentials) -> None:
    try:
        pass
    except HttpError as e:
        print(f"⚠️ Failed to apply heading/text styles: {e}")


def upload_and_convert_docx_to_gdoc(local_docx: str, title: str, folder_id: str, creds: Credentials) -> str:
    drive = build('drive', 'v3', credentials=creds)
    media = MediaFileUpload(local_docx, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
    metadata = {'name': title, 'parents': [folder_id], 'mimeType': 'application/vnd.google-apps.document'}
    max_retries = int(os.getenv('API_MAX_RETRIES', '6'))
    backoff = float(os.getenv('API_BACKOFF_START', '1.0'))
    backoff_cap = float(os.getenv('API_BACKOFF_CAP', '8.0'))
    for attempt in range(max_retries):
        try:
            file = drive.files().create(body=metadata, media_body=media, fields='id').execute()
            return file['id']
        except HttpError as e:
            status = getattr(getattr(e, 'resp', None), 'status', None)
            if status in (429, 403, 500, 503):
                time.sleep(backoff + random.uniform(0, 0.25))
                backoff = min(backoff_cap, backoff * 2.0)
                continue
            raise

def create_google_doc(local_docx: str, title: str, folder_id: str, creds: Credentials) -> str:
    doc_id = upload_and_convert_docx_to_gdoc(local_docx, title, folder_id, creds)
    apply_named_styles(local_docx, doc_id, creds)
    return doc_id
